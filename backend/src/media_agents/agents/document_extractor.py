"""Extract text content from base64-encoded document data URLs.

Supports PDF (via pypdf), DOCX (via python-docx), and plain text / markdown.
On any failure, returns a descriptive fallback string rather than raising.
"""

from __future__ import annotations

import base64
import io
import logging

logger = logging.getLogger(__name__)

_MAX_TEXT_LENGTH = 50_000


def _decode_data_url(data_url: str) -> bytes:
    """Strip the ``data:<mime>;base64,`` prefix and decode."""
    if "," not in data_url:
        raise ValueError("Invalid data URL: missing comma separator")
    _, _, encoded = data_url.partition(",")
    return base64.b64decode(encoded)


def _extract_pdf(raw: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(raw))
    pages: list[str] = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n\n".join(pages)


def _extract_docx(raw: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(raw))
    parts: list[str] = []
    
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    stripped = p.text.strip()
                    if stripped:
                        parts.append(stripped)
                        
    return "\n\n".join(parts)


def _truncate(text: str) -> str:
    if len(text) <= _MAX_TEXT_LENGTH:
        return text
    suffix = "…[truncated]"
    limit = max(0, _MAX_TEXT_LENGTH - len(suffix))
    return text[:limit] + suffix


def extract_text(data_url: str, mime_type: str, filename: str) -> str:
    """Decode a base64 data URL and extract text content.

    Returns extracted text, or a fallback string on failure.
    """
    mime_type = mime_type.split(";")[0].strip().lower()

    try:
        raw = _decode_data_url(data_url)
    except Exception as e:
        logger.warning("Failed to decode base64 for %s: %s", filename, e)
        return f"[Could not extract text from {filename}]"

    try:
        if mime_type == "application/pdf":
            text = _extract_pdf(raw)
        elif mime_type in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ):
            text = _extract_docx(raw)
        elif mime_type in ("text/plain", "text/markdown"):
            text = raw.decode("utf-8", errors="replace")
        else:
            logger.warning("Unsupported mime type %s for file %s", mime_type, filename)
            return f"[Could not extract text from {filename}]"
    except Exception as e:
        logger.warning("Text extraction failed for %s (%s): %s", filename, mime_type, e)
        return f"[Could not extract text from {filename}]"

    return _truncate(text)
